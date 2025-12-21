# ==============================================================
# STREAMLIT DASHBOARD + EXPORT SENTIMENT ANALYSIS (ML + BERT)
# Author: Sparshi Jain
# ==============================================================

# ==============================================================
# STREAMLIT DASHBOARD + EXPORT SENTIMENT ANALYSIS (ML + BERT)
# Author: Sparshi Jain
# ==============================================================

import streamlit as st
import pandas as pd
import pickle
import matplotlib.pyplot as plt
import os
import requests
import torch

from transformers import AutoTokenizer, AutoModelForSequenceClassification

# Optional exports (safe)
try:
    from docx import Document
except ImportError:
    Document = None

try:
    from fpdf import FPDF
except ImportError:
    FPDF = None


# ==============================================================
# MODEL DOWNLOAD URLS (GitHub RAW)
# ==============================================================

TFIDF_URL = "https://raw.githubusercontent.com/sparshi15/Research-Paper-Abstract-Sentiment-Analysis-By-Machine-Learning-/main/tfidf_vectorizer.pkl"
LR_URL = "https://raw.githubusercontent.com/sparshi15/Research-Paper-Abstract-Sentiment-Analysis-By-Machine-Learning-/main/logistic_regression_model.pkl"



# ==============================================================
# DOWNLOAD HELPER
# ==============================================================

def download_model(url, filename):
    if not os.path.exists(filename):
        with st.spinner(f"Downloading {filename}..."):
            r = requests.get(url)
            r.raise_for_status()
            with open(filename, "wb") as f:
                f.write(r.content)


# ==============================================================
# LOAD ML MODELS (TF-IDF)
# ==============================================================

@st.cache_resource
def load_ml_models():
    download_model(TFIDF_URL, "tfidf_vectorizer.pkl")
    download_model(LR_URL, "logistic_regression_model.pkl")
    

    tfidf = pickle.load(open("tfidf_vectorizer.pkl", "rb"))
    lr_model = pickle.load(open("logistic_regression_model.pkl", "rb"))
   

    return tfidf, lr_model


# ==============================================================
# LOAD BERT MODEL
# ==============================================================

@st.cache_resource
def load_bert():
    tokenizer = AutoTokenizer.from_pretrained(
        "nlptown/bert-base-multilingual-uncased-sentiment"
    )
    model = AutoModelForSequenceClassification.from_pretrained(
        "nlptown/bert-base-multilingual-uncased-sentiment"
    )
    return tokenizer, model


tfidf, lr_model, rf_model = load_ml_models()
bert_tokenizer, bert_model = load_bert()


# ==============================================================
# PREDICTION FUNCTIONS
# ==============================================================

def predict_tf_idf(text, model_name):
    X = tfidf.transform([text])
    if model_name == "Logistic Regression (TF-IDF)":
        return lr_model.predict(X)[0]
    else:
        return rf_model.predict(X)[0]


def predict_bert(text):
    tokens = bert_tokenizer(
        text, return_tensors="pt", truncation=True, padding=True
    )
    with torch.no_grad():
        output = bert_model(**tokens)

    pred = torch.argmax(output.logits).item()
    return ["negative", "negative", "neutral", "positive", "positive"][pred]


# ==============================================================
# EXPORT FUNCTIONS
# ==============================================================

def export_to_word(df):
    if Document is None:
        st.error("python-docx not installed.")
        return None

    doc = Document()
    doc.add_heading("Sentiment Analysis Report", level=1)

    table = doc.add_table(rows=1, cols=len(df.columns))
    table.style = "Table Grid"

    for i, col in enumerate(df.columns):
        table.rows[0].cells[i].text = col

    for _, row in df.iterrows():
        row_cells = table.add_row().cells
        for i, val in enumerate(row):
            row_cells[i].text = str(val)

    filename = "Sentiment_Report.docx"
    doc.save(filename)
    return filename


def export_to_pdf(df):
    if FPDF is None:
        st.error("fpdf not installed.")
        return None

    pdf = FPDF()
    pdf.add_page()
    pdf.set_font("Arial", size=9)
    pdf.cell(200, 10, "Sentiment Analysis Report", ln=True, align="C")

    for _, row in df.iterrows():
        pdf.multi_cell(0, 5, str(row.to_dict()))

    filename = "Sentiment_Report.pdf"
    pdf.output(filename)
    return filename


# ==============================================================
# STREAMLIT UI
# ==============================================================

st.set_page_config(page_title="Research Sentiment Dashboard", layout="wide")
st.title("📊 Research Abstract Sentiment Dashboard")

st.sidebar.title("⚙ Settings")

model_choice = st.sidebar.selectbox(
    "Choose Sentiment Model",
    [
        "Logistic Regression (TF-IDF)",
        "BERT (Pretrained)"
    ]
)

mode = st.sidebar.radio(
    "Mode",
    ["Upload CSV", "Manual Text Prediction"]
)

st.sidebar.markdown("---")


# ==============================================================
# CSV UPLOAD MODE
# ==============================================================

if mode == "Upload CSV":
    uploaded_file = st.file_uploader(
        "📂 Upload CSV (must include 'abstract' column)",
        type=["csv"]
    )

    if uploaded_file:
        df = pd.read_csv(uploaded_file)

        if "abstract" not in df.columns:
            st.error("❌ CSV must contain an 'abstract' column.")
        else:
            if st.button("🚀 Predict Sentiment"):
                with st.spinner("Analyzing abstracts..."):
                    if model_choice == "BERT (Pretrained)":
                        df["Predicted Sentiment"] = df["abstract"].apply(predict_bert)
                    else:
                        df["Predicted Sentiment"] = df["abstract"].apply(
                            lambda x: predict_tf_idf(x, model_choice)
                        )

                st.subheader("✅ Prediction Results")
                st.dataframe(df.head())

                st.subheader("📌 Metrics Summary")
                sentiment_counts = df["Predicted Sentiment"].value_counts(normalize=True) * 100

                col1, col2, col3 = st.columns(3)
                col1.metric("Positive %", f"{sentiment_counts.get('positive', 0):.2f}%")
                col2.metric("Neutral %", f"{sentiment_counts.get('neutral', 0):.2f}%")
                col3.metric("Negative %", f"{sentiment_counts.get('negative', 0):.2f}%")

                st.subheader("🟢 Sentiment Distribution")
                fig, ax = plt.subplots()
                df["Predicted Sentiment"].value_counts().plot(
                    kind="pie", autopct="%1.1f%%", ax=ax
                )
                ax.set_ylabel("")
                st.pyplot(fig)

                st.subheader("⬇ Export Results")
                st.download_button(
                    "Download CSV",
                    df.to_csv(index=False),
                    "sentiment_results.csv"
                )

                if st.button("Download Word Report"):
                    file = export_to_word(df)
                    if file:
                        st.download_button("Download Word", open(file, "rb"), file)

                if st.button("Download PDF Report"):
                    file = export_to_pdf(df)
                    if file:
                        st.download_button("Download PDF", open(file, "rb"), file)


# ==============================================================
# MANUAL TEXT MODE
# ==============================================================

else:
    text = st.text_area("✍ Enter Research Abstract:")

    if st.button("Predict"):
        if not text.strip():
            st.warning("Please enter some text.")
        else:
            result = (
                predict_bert(text)
                if model_choice == "BERT (Pretrained)"
                else predict_tf_idf(text, model_choice)
            )
            st.success(f"✅ Predicted Sentiment: **{result.upper()}**")




